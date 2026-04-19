"""
SECURE ACADEMIC DOCUMENT VERIFICATION SYSTEM
=============================================
Main Flask Application

Cybersecurity & Blockchain Concepts Implemented:
  ✔ SHA-256 Document Hashing
  ✔ RSA-2048 Digital Signatures (Asymmetric Cryptography)
  ✔ AES-256-CBC File Encryption (Symmetric Cryptography)
  ✔ HMAC-SHA256 Message Authentication Code
  ✔ Blockchain with Proof of Work
  ✔ Merkle Tree for transaction integrity
  ✔ Chain Validation (Tamper Detection)
  ✔ Base64 Encoding
  ✔ Secure File Handling
"""

import os
import json
import base64
import hashlib
import hmac
import difflib
import re
import uuid
import unicodedata
import tempfile
from datetime import datetime, timezone
from functools import wraps
from pathlib import Path

from flask import (Flask, render_template, request, jsonify,
                   redirect, url_for, flash, send_file, session)
from werkzeug.utils import secure_filename

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from rapidocr_onnxruntime import RapidOCR
except ImportError:
    RapidOCR = None

try:
    import cv2
except ImportError:
    cv2 = None
try:
    import numpy as np
except ImportError:
    np = None
try:
    import pypdfium2 as pdfium
except ImportError:
    pdfium = None

from blockchain import Blockchain
from crypto_utils import (
    sha256_file, sha256_string, sha256_bytes,
    generate_hmac, verify_hmac,
    sign_hash, verify_rsa_signature,
    aes_encrypt, aes_decrypt,
    load_or_create_keys, build_document_record,
    load_or_create_certificate, sign_with_certificate,
    verify_certificate_signature, certificate_summary,
)

# ─────────────────────────────────────────────────────────────
# APP SETUP
# ─────────────────────────────────────────────────────────────

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET", "SADV_SECRET_KEY_2024")

UPLOAD_FOLDER      = "uploads"
ORIGINALS_DIR      = os.path.join(UPLOAD_FOLDER, "originals")
ENCRYPTED_DIR      = os.path.join(UPLOAD_FOLDER, "encrypted")
ALLOWED_EXTENSIONS = {"pdf", "doc", "docx", "png", "jpg", "jpeg", "txt"}
USERS_FILE         = "users.json"
AUDIT_LOG_FILE     = "audit_log.json"
DOCUMENT_STATE_FILE = "document_state.json"

DEFAULT_USERS = [
    {"username": "admin", "role": "Admin", "password": "Admin@123"},
    {"username": "verifier", "role": "Verifier", "password": "Verifier@123"},
]

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(ORIGINALS_DIR, exist_ok=True)
os.makedirs(ENCRYPTED_DIR, exist_ok=True)
os.makedirs("keys", exist_ok=True)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024  # 16 MB

# Initialise blockchain and cryptographic keys once at startup
blockchain = Blockchain(difficulty=2)
private_key_pem, public_key_pem, aes_key = load_or_create_keys()
institution_certificate_pem = load_or_create_certificate(private_key_pem, public_key_pem)
institution_certificate_info = certificate_summary(institution_certificate_pem)
ocr_engine = RapidOCR() if RapidOCR else None
qr_detector = cv2.QRCodeDetector() if cv2 else None
barcode_detector = cv2.barcode_BarcodeDetector() if cv2 and hasattr(cv2, "barcode_BarcodeDetector") else None

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def allowed_file(filename: str) -> bool:
    return "." in filename and \
           filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def generate_name_variants(filename: str) -> list[str]:
    """
    Normalize filename stems to improve tampered-file matching.
    Example: marksheet_tampered_v2 -> marksheet
    """
    stem = Path(filename or "").stem.lower().strip()
    if not stem:
        return []
    variants = {stem}
    cleaned = re.sub(r"[\W_]+", " ", stem).strip()
    if cleaned:
        variants.add(cleaned.replace(" ", "_"))

    removable_tokens = {
        "tampered", "modified", "change", "changed", "edit", "edited",
        "fake", "copy", "new", "final", "updated", "v2", "v3", "v4",
        "rev", "revised", "candidate", "upload", "verified",
    }
    tokens = [t for t in re.split(r"[\W_]+", stem) if t]
    core_tokens = [t for t in tokens if t not in removable_tokens and not t.startswith("v")]
    if core_tokens:
        variants.add("_".join(core_tokens))
        variants.add("".join(core_tokens))
    return [v for v in variants if v]


def file_size_kb(path: str) -> float:
    return round(os.path.getsize(path) / 1024, 2)


def normalize_text_for_hash(text: str) -> str:
    """Canonical representation focused on document meaning, not formatting."""
    if not text:
        return ""
    normalized = unicodedata.normalize("NFKC", text)
    normalized = normalized.replace("\u00a0", " ")
    normalized = re.sub(r"[\u200b-\u200d\ufeff]", "", normalized)
    normalized = normalized.lower()
    # Preserve meaningful inline symbols used in grades, IDs, and amounts.
    tokens = re.findall(r"[a-z0-9]+(?:[+./:-][a-z0-9]+)*", normalized)
    return " ".join(tokens)


def compute_text_hash_from_text(text: str) -> str:
    normalized = normalize_text_for_hash(text)
    return sha256_string(normalized) if normalized else ""


def compute_text_hash_from_file(file_path: str) -> str:
    text = extract_text(file_path)
    return compute_text_hash_from_text(text)


def load_document_state() -> dict:
    if os.path.exists(DOCUMENT_STATE_FILE):
        try:
            with open(DOCUMENT_STATE_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"revoked_hashes": []}


def save_document_state(data: dict) -> None:
    with open(DOCUMENT_STATE_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


def is_document_active(doc_hash: str) -> bool:
    state = load_document_state()
    return doc_hash not in set(state.get("revoked_hashes", []))


def revoke_document(doc_hash: str) -> None:
    state = load_document_state()
    revoked = set(state.get("revoked_hashes", []))
    revoked.add(doc_hash)
    state["revoked_hashes"] = sorted(revoked)
    save_document_state(state)


def restore_document(doc_hash: str) -> None:
    state = load_document_state()
    revoked = set(state.get("revoked_hashes", []))
    if doc_hash in revoked:
        revoked.remove(doc_hash)
    state["revoked_hashes"] = sorted(revoked)
    save_document_state(state)


def get_registered_documents_for_verifier() -> list[dict]:
    """
    Build dropdown options of registered originals for verifier flow.
    """
    docs = []
    seen_hashes = set()
    for block_dict, tx in blockchain.get_all_document_records():
        doc_hash = tx.get("document_hash")
        if not doc_hash or doc_hash in seen_hashes:
            continue
        if not is_document_active(doc_hash):
            continue
        seen_hashes.add(doc_hash)
        docs.append({
            "hash": doc_hash,
            "name": tx.get("document_name", "unknown"),
            "category": tx.get("document_category", "Unknown"),
            "student_name": tx.get("student_name", "-"),
            "degree": tx.get("degree", "-"),
            "institution": tx.get("institution", "-"),
            "issued_at": tx.get("issued_at", ""),
            "block_index": block_dict.get("index"),
        })
    return docs


def get_registered_documents_for_admin() -> list[dict]:
    docs = []
    seen_hashes = set()
    revoked = set(load_document_state().get("revoked_hashes", []))
    for block_dict, tx in blockchain.get_all_document_records():
        doc_hash = tx.get("document_hash")
        if not doc_hash or doc_hash in seen_hashes:
            continue
        seen_hashes.add(doc_hash)
        docs.append({
            "hash": doc_hash,
            "name": tx.get("document_name", "unknown"),
            "category": tx.get("document_category", "Unknown"),
            "student_name": tx.get("student_name", "-"),
            "degree": tx.get("degree", "-"),
            "institution": tx.get("institution", "-"),
            "issued_at": tx.get("issued_at", ""),
            "block_index": block_dict.get("index"),
            "active": doc_hash not in revoked,
        })
    return docs


def hash_password(password: str, salt: bytes | None = None) -> tuple[str, str]:
    salt = salt or os.urandom(16)
    digest = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt,
        150000,
        dklen=32,
    )
    return base64.b64encode(digest).decode(), base64.b64encode(salt).decode()


def verify_password(password: str, salt_b64: str, expected_hash: str) -> bool:
    salt = base64.b64decode(salt_b64)
    digest = hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt,
        150000,
        dklen=32,
    )
    return hmac.compare_digest(base64.b64encode(digest).decode(), expected_hash)


def load_user_db() -> dict:
    if os.path.exists(USERS_FILE):
        try:
            with open(USERS_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"users": []}


def save_user_db(data: dict) -> None:
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2)


def ensure_user_db() -> None:
    data = load_user_db()
    if not data.get("users"):
        data["users"] = []
    existing_usernames = {u["username"] for u in data["users"]}
    for template in DEFAULT_USERS:
        if template["username"] not in existing_usernames:
            password_hash, salt = hash_password(template["password"])
            data["users"].append({
                "username": template["username"],
                "role": template["role"],
                "password_hash": password_hash,
                "salt": salt,
            })
    save_user_db(data)


def get_user(username: str) -> dict | None:
    data = load_user_db()
    for user in data.get("users", []):
        if user.get("username") == username:
            return user
    return None


def _canonical_audit_entry(entry: dict) -> dict:
    return {
        "entry_id": entry.get("entry_id"),
        "timestamp": entry.get("timestamp"),
        "user": entry.get("user"),
        "role": entry.get("role"),
        "action": entry.get("action"),
        "document_name": entry.get("document_name"),
        "status": entry.get("status"),
        "details": entry.get("details"),
        "deleted": entry.get("deleted", False),
        "deleted_at": entry.get("deleted_at"),
        "deleted_by": entry.get("deleted_by"),
        "previous_entry_hash": entry.get("previous_entry_hash", ""),
    }


def _rehash_audit_entries(logs: list[dict]) -> list[dict]:
    previous_hash = ""
    for entry in logs:
        entry.setdefault("entry_id", str(uuid.uuid4()))
        entry["previous_entry_hash"] = previous_hash
        entry["entry_hash"] = sha256_string(json.dumps(_canonical_audit_entry(entry), sort_keys=True))
        previous_hash = entry["entry_hash"]
    return logs


def load_audit_log() -> list[dict]:
    logs = []
    if os.path.exists(AUDIT_LOG_FILE):
        try:
            with open(AUDIT_LOG_FILE, "r", encoding="utf-8") as f:
                logs = json.load(f)
        except Exception:
            logs = []
    logs = _rehash_audit_entries(logs)
    with open(AUDIT_LOG_FILE, "w", encoding="utf-8") as f:
        json.dump(logs, f, indent=2)
    return logs


def save_audit_log(logs: list[dict]) -> None:
    logs = _rehash_audit_entries(logs)
    with open(AUDIT_LOG_FILE, "w", encoding="utf-8") as f:
        json.dump(logs, f, indent=2)


def verify_audit_chain(logs: list[dict]) -> bool:
    previous_hash = ""
    for entry in logs:
        expected = sha256_string(json.dumps({
            **_canonical_audit_entry(entry),
            "previous_entry_hash": previous_hash,
        }, sort_keys=True))
        if entry.get("previous_entry_hash", "") != previous_hash:
            return False
        if entry.get("entry_hash") != expected:
            return False
        previous_hash = entry.get("entry_hash", "")
    return True


def append_audit_entry(action: str, status: str, document_name: str | None = None, details: str | None = None, username: str | None = None, role: str | None = None) -> None:
    entry = {
        "entry_id": str(uuid.uuid4()),
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "user": username or session.get("username"),
        "role": role or session.get("role"),
        "action": action,
        "document_name": document_name,
        "status": status,
        "details": details,
        "deleted": False,
    }
    logs = load_audit_log()
    logs.append(entry)
    save_audit_log(logs)


def get_audit_entries(limit: int = 50, role: str | None = None, username: str | None = None) -> list:
    logs = load_audit_log()
    filtered = []
    for entry in reversed(logs):
        if entry.get("deleted"):
            continue
        if role and entry.get("role") != role:
            continue
        if username and entry.get("user") != username:
            continue
        filtered.append(entry)
    return filtered[:limit]


def delete_audit_entry(entry_id: str, requester_role: str, requester_user: str) -> bool:
    logs = load_audit_log()

    index_to_remove = None
    for idx, entry in enumerate(logs):
        if entry.get("entry_id") != entry_id:
            continue
        if requester_role == "Admin":
            index_to_remove = idx
            break
        if requester_role == "Verifier" and entry.get("user") == requester_user and entry.get("action", "").startswith("verify_"):
            index_to_remove = idx
            break

    if index_to_remove is None:
        return False
    logs[index_to_remove]["deleted"] = True
    logs[index_to_remove]["deleted_at"] = datetime.now(timezone.utc).isoformat()
    logs[index_to_remove]["deleted_by"] = requester_user
    save_audit_log(logs)
    return True


def format_timestamp(value: str) -> str:
    try:
        return datetime.fromisoformat(value).strftime("%Y-%m-%d %H:%M:%S UTC")
    except Exception:
        return value


def extract_metadata(file_path: str) -> dict:
    metadata = {
        "file_name": os.path.basename(file_path),
        "file_size": os.path.getsize(file_path),
        "page_count": None,
        "author": None,
        "created_at": None,
        "modified_at": None,
    }
    suffix = Path(file_path).suffix.lower()
    if suffix == ".pdf" and PdfReader:
        try:
            reader = PdfReader(file_path)
            metadata["page_count"] = len(reader.pages)
            info = reader.metadata
            if info:
                metadata["author"] = getattr(info, "author", None) or info.get("/Author")
                created = getattr(info, "creation_date", None) or info.get("/CreationDate")
                modified = getattr(info, "mod_date", None) or info.get("/ModDate")
                if created:
                    metadata["created_at"] = str(created)
                if modified:
                    metadata["modified_at"] = str(modified)
        except Exception:
            pass
    elif suffix == ".docx" and Document:
        try:
            doc = Document(file_path)
            metadata["page_count"] = max(len(doc.paragraphs), 1)
            props = doc.core_properties
            metadata["author"] = props.author
            if props.created:
                metadata["created_at"] = props.created.isoformat()
            if props.modified:
                metadata["modified_at"] = props.modified.isoformat()
        except Exception:
            pass
    elif suffix == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                metadata["page_count"] = max(len(f.read().splitlines()), 1)
        except Exception:
            pass
    else:
        try:
            stats = os.stat(file_path)
            metadata["created_at"] = datetime.fromtimestamp(stats.st_ctime, timezone.utc).isoformat()
            metadata["modified_at"] = datetime.fromtimestamp(stats.st_mtime, timezone.utc).isoformat()
        except Exception:
            pass
    return metadata


def extract_visual_codes_from_image(image_bgr) -> list[str]:
    payloads = []
    if image_bgr is None or cv2 is None:
        return payloads

    if qr_detector:
        try:
            ok, decoded_info, _points, _ = qr_detector.detectAndDecodeMulti(image_bgr)
            if ok and decoded_info:
                payloads.extend([item.strip() for item in decoded_info if item and item.strip()])
        except Exception:
            pass
        if not payloads:
            try:
                payload, _points, _ = qr_detector.detectAndDecode(image_bgr)
                if payload and payload.strip():
                    payloads.append(payload.strip())
            except Exception:
                pass

    if barcode_detector:
        try:
            decoded = barcode_detector.detectAndDecode(image_bgr)
            if isinstance(decoded, tuple):
                for item in decoded:
                    if isinstance(item, (list, tuple)):
                        for sub in item:
                            if isinstance(sub, str) and sub.strip():
                                payloads.append(sub.strip())
                    elif isinstance(item, str) and item.strip():
                        payloads.append(item.strip())
        except Exception:
            pass

    return list(dict.fromkeys(payloads))


def extract_text_from_image(image_bgr) -> list[str]:
    lines = []
    if ocr_engine and image_bgr is not None:
        try:
            result, _ = ocr_engine(image_bgr)
            if result:
                for item in result:
                    if len(item) >= 2 and item[1]:
                        lines.append(str(item[1]).strip())
        except Exception:
            pass
    for payload in extract_visual_codes_from_image(image_bgr):
        lines.append(f"CODE_PAYLOAD: {payload}")
    return [line for line in lines if line]


def render_pdf_pages(file_path: str, scale: float = 2.0) -> list:
    if not pdfium or not np or not cv2:
        return []
    try:
        document = pdfium.PdfDocument(file_path)
        pages = []
        for index in range(len(document)):
            page = document[index]
            rendered = page.render(scale=scale).to_pil()
            pages.append(cv2.cvtColor(np.array(rendered), cv2.COLOR_RGB2BGR))
        return pages
    except Exception:
        return []


def extract_text(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""
    if suffix == ".pdf":
        extracted_parts = []
        if PdfReader:
            try:
                reader = PdfReader(file_path)
                for page in reader.pages:
                    extracted_parts.append(page.extract_text() or "")
            except Exception:
                pass
        # OCR rendered pages so scanned PDFs participate in content hashing.
        for page_image in render_pdf_pages(file_path):
            extracted_parts.extend(extract_text_from_image(page_image))
        return "\n".join(part for part in extracted_parts if part and str(part).strip())
    if suffix == ".docx" and Document:
        try:
            doc = Document(file_path)
            lines = [p.text for p in doc.paragraphs if p.text]
            rel_parts = list(doc.part.related_parts.values())
            for related in rel_parts:
                try:
                    if not getattr(related, "content_type", "").startswith("image/"):
                        continue
                    if not cv2 or not np:
                        continue
                    image_bytes = related.blob
                    image_array = np.frombuffer(image_bytes, dtype=np.uint8)
                    image_bgr = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
                    lines.extend(extract_text_from_image(image_bgr))
                except Exception:
                    continue
            return "\n".join(line for line in lines if line)
        except Exception:
            return ""
    if suffix in {".png", ".jpg", ".jpeg"} and cv2:
        try:
            image_bgr = cv2.imread(file_path)
            return "\n".join(extract_text_from_image(image_bgr))
        except Exception:
            return ""
    return ""


def compare_metadata(stored_meta: dict, uploaded_meta: dict) -> tuple[bool, dict]:
    changes = {}
    # Volatile fields like timestamps and raw file size change across formats
    # even when the underlying text is the same, so they are not used as tamper signals.
    for key in ["author", "page_count"]:
        stored_value = stored_meta.get(key)
        uploaded_value = uploaded_meta.get(key)
        if stored_value != uploaded_value:
            changes[key] = {"stored": stored_value, "uploaded": uploaded_value}
    return bool(changes), changes


def compute_content_change(original_text: str, uploaded_text: str) -> tuple[float, int, int]:
    if not original_text and not uploaded_text:
        return 0.0, 0, 0
    matcher = difflib.SequenceMatcher(None, original_text, uploaded_text)
    total_length = max(len(original_text), len(uploaded_text), 1)
    matching_length = sum(block.size for block in matcher.get_matching_blocks())
    diff_length = max(total_length - matching_length, 0)
    return round((diff_length / total_length) * 100.0, 1), diff_length, total_length


def compute_structure_change(stored_meta: dict, uploaded_meta: dict, content_equivalent: bool = False) -> dict:
    if content_equivalent:
        return {
            "structure_change_pct": 0.0,
            "file_size_change_pct": 0.0,
            "page_count_change_pct": 0.0,
        }

    original_size = int(stored_meta.get("file_size") or 0)
    uploaded_size = int(uploaded_meta.get("file_size") or 0)
    if original_size <= 0:
        size_change_pct = 100.0 if uploaded_size > 0 else 0.0
    else:
        size_change_pct = round(
            min(abs(original_size - uploaded_size) / max(original_size, uploaded_size) * 100.0, 100.0),
            1,
        )

    original_pages = stored_meta.get("page_count")
    uploaded_pages = uploaded_meta.get("page_count")
    page_change_pct = None
    if isinstance(original_pages, int) and isinstance(uploaded_pages, int):
        if original_pages <= 0:
            page_change_pct = 100.0 if uploaded_pages > 0 else 0.0
        else:
            page_change_pct = round(
                min(abs(original_pages - uploaded_pages) / max(original_pages, uploaded_pages) * 100.0, 100.0),
                1,
            )

    components = [size_change_pct]
    if page_change_pct is not None:
        components.append(page_change_pct)
    structure_change_pct = round(sum(components) / len(components), 1)
    return {
        "structure_change_pct": structure_change_pct,
        "file_size_change_pct": size_change_pct,
        "page_count_change_pct": page_change_pct,
    }


def classify_risk(score: float) -> str:
    if score <= 20:
        return "Safe"
    if score <= 50:
        return "Suspicious"
    if score <= 80:
        return "High Risk"
    return "Fake"


def generate_text_diff(original_text: str, uploaded_text: str) -> str:
    original_lines = original_text.splitlines()
    uploaded_lines = uploaded_text.splitlines()
    diff = list(difflib.unified_diff(
        original_lines,
        uploaded_lines,
        fromfile="Original",
        tofile="Uploaded",
        lineterm="",
    ))
    if not diff:
        return "No textual differences detected or text extraction was unavailable."
    return "\n".join(diff)


def detect_text_changes(original_text: str, uploaded_text: str, max_items: int = 25) -> dict:
    original_lines = original_text.splitlines()
    uploaded_lines = uploaded_text.splitlines()
    matcher = difflib.SequenceMatcher(None, original_lines, uploaded_lines)
    added = []
    removed = []
    modified = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "insert":
            segment = "\n".join(uploaded_lines[j1:j2]).strip()
            if segment:
                added.append({"line_start": j1 + 1, "text": segment[:800]})
        elif tag == "delete":
            segment = "\n".join(original_lines[i1:i2]).strip()
            if segment:
                removed.append({"line_start": i1 + 1, "text": segment[:800]})
        elif tag == "replace":
            original_segment = "\n".join(original_lines[i1:i2]).strip()
            uploaded_segment = "\n".join(uploaded_lines[j1:j2]).strip()
            if original_segment or uploaded_segment:
                modified.append({
                    "original_line_start": i1 + 1,
                    "uploaded_line_start": j1 + 1,
                    "original_text": original_segment[:800],
                    "uploaded_text": uploaded_segment[:800],
                })

    return {
        "added": added[:max_items],
        "removed": removed[:max_items],
        "modified": modified[:max_items],
        "totals": {
            "added": len(added),
            "removed": len(removed),
            "modified": len(modified),
        },
    }


def extract_code_payloads_from_text(text: str) -> list[str]:
    payloads = []
    for line in (text or "").splitlines():
        if line.startswith("CODE_PAYLOAD:"):
            payload = line.split("CODE_PAYLOAD:", 1)[1].strip()
            if payload:
                payloads.append(payload)
    return list(dict.fromkeys(payloads))


def compute_similarity_pct(original_text: str, uploaded_text: str) -> float:
    if not original_text and not uploaded_text:
        return 0.0
    return round(difflib.SequenceMatcher(None, original_text, uploaded_text).ratio() * 100.0, 1)


def build_signal(name: str, risk: float | None, weight: float, status: str, reason: str) -> dict:
    return {
        "name": name,
        "risk": risk,
        "weight": weight,
        "status": status,
        "reason": reason,
        "available": risk is not None,
    }


def compute_ensemble_verdict(signals: list[dict]) -> tuple[float, str]:
    weighted_total = 0.0
    total_weight = 0.0
    for signal in signals:
        if signal.get("available") and signal.get("risk") is not None:
            weighted_total += float(signal["risk"]) * float(signal["weight"])
            total_weight += float(signal["weight"])
    risk_score = round(weighted_total / total_weight, 1) if total_weight else 100.0
    return risk_score, classify_risk(risk_score)


def build_analysis_reasons(
    *,
    content_equivalent: bool,
    original_text_available: bool,
    uploaded_text_available: bool,
    content_change_pct: float,
    metadata_changed: bool,
    structure_change_pct: float,
    file_size_change_pct: float,
    classification: str,
    signal_results: list[dict] | None = None,
) -> list[str]:
    reasons = []

    if content_equivalent:
        reasons.append("Extracted text matches the registered original after normalization.")

    if not original_text_available or not uploaded_text_available:
        reasons.append("Text extraction was unavailable for one or both files, so the result relies on binary and structural checks.")

    if content_change_pct > 0:
        reasons.append(f"Document content differs from the registered original by {content_change_pct}%.")

    if metadata_changed and not content_equivalent:
        reasons.append("Stable metadata fields differ from the registered original.")

    if structure_change_pct > 0:
        reasons.append(f"File structure differs by {structure_change_pct}% and file size differs by {file_size_change_pct}%.")

    for signal in signal_results or []:
        if signal.get("available"):
            reasons.append(f"{signal['name']}: {signal['reason']}")

    if classification == "Safe" and not content_equivalent:
        reasons.append("The mismatch indicators are low, so the document is not strongly flagged as tampered.")
    elif classification == "Suspicious":
        reasons.append("There are moderate differences, so the document needs manual review.")
    elif classification == "Fake":
        reasons.append("The detected differences are severe enough to strongly suggest tampering or substitution.")

    if not reasons:
        reasons.append("No suspicious differences were detected.")

    return reasons


def analyze_uploaded_against_stored(stored_record: dict, uploaded_path: str) -> dict:
    stored_meta = stored_record.get("metadata", {})
    uploaded_meta = extract_metadata(uploaded_path)

    original_text = extract_text(stored_record.get("original_file_path", ""))
    uploaded_text = extract_text(uploaded_path)
    normalized_original_text = normalize_text_for_hash(original_text)
    normalized_uploaded_text = normalize_text_for_hash(uploaded_text)
    text_equivalent = bool(
        normalized_original_text and
        normalized_uploaded_text and
        normalized_original_text == normalized_uploaded_text
    )
    metadata_changed, metadata_diff = compare_metadata(stored_meta, uploaded_meta)
    content_change_pct, diff_length, total_length = compute_content_change(
        normalized_original_text,
        normalized_uploaded_text,
    )
    similarity_pct = compute_similarity_pct(normalized_original_text, normalized_uploaded_text)
    original_codes = extract_code_payloads_from_text(original_text)
    uploaded_codes = extract_code_payloads_from_text(uploaded_text)
    code_equivalent = original_codes == uploaded_codes if (original_codes or uploaded_codes) else True
    content_equivalent = text_equivalent and code_equivalent

    structure = compute_structure_change(stored_meta, uploaded_meta, content_equivalent=content_equivalent)
    structure_change_pct = structure["structure_change_pct"]
    metadata_score = 100.0 if metadata_changed and not content_equivalent else 0.0
    if original_codes or uploaded_codes:
        if original_codes == uploaded_codes and original_codes:
            code_risk = 0.0
            code_status = "matched"
            code_reason = "Embedded QR/barcode payloads match."
        elif original_codes and uploaded_codes:
            code_risk = 100.0
            code_status = "mismatch"
            code_reason = "Embedded QR/barcode payloads differ."
        else:
            code_risk = 70.0
            code_status = "missing"
            code_reason = "A QR/barcode payload is missing from one of the compared files."
    else:
        code_risk = None
        code_status = "unavailable"
        code_reason = "No QR/barcode payload was detected."

    signal_results = [
        build_signal(
            "Text content",
            0.0 if text_equivalent else max(content_change_pct * 12.0, 35.0) if normalized_original_text and normalized_uploaded_text else None,
            0.35,
            "matched" if text_equivalent else "changed",
            "Normalized extracted text matches." if text_equivalent else f"Normalized extracted text similarity is {similarity_pct}%.",
        ),
        build_signal(
            "Content hash",
            0.0 if text_equivalent else 100.0 if normalized_original_text and normalized_uploaded_text else None,
            0.15,
            "matched" if text_equivalent else "mismatch",
            "Text-derived content hash matches." if text_equivalent else "Text-derived content hash does not match.",
        ),
        build_signal("QR / barcode", code_risk, 0.20, code_status, code_reason),
        build_signal(
            "Metadata",
            100.0 if metadata_changed and not content_equivalent else 0.0,
            0.10,
            "changed" if metadata_changed and not content_equivalent else "matched",
            "Stable metadata differs." if metadata_changed and not content_equivalent else "Stable metadata matches.",
        ),
        build_signal(
            "Structure",
            structure_change_pct,
            0.10,
            "changed" if structure_change_pct > 0 else "matched",
            f"Structure change measured at {structure_change_pct}%.",
        ),
        build_signal(
            "Binary hash",
            100.0,
            0.10,
            "mismatch",
            "Raw file bytes differ, so the exact binary file is not identical.",
        ),
    ]

    if content_equivalent:
        risk_score = 0.0
        classification = "Safe"
    else:
        risk_score, classification = compute_ensemble_verdict(signal_results)
    if (
        not content_equivalent
        and bool(original_text.strip())
        and bool(uploaded_text.strip())
        and content_change_pct > 0
        and classification == "Safe"
    ):
        classification = "Suspicious"
        risk_score = max(risk_score, 25.0)
    reasons = build_analysis_reasons(
        content_equivalent=content_equivalent,
        original_text_available=bool(original_text.strip()),
        uploaded_text_available=bool(uploaded_text.strip()),
        content_change_pct=content_change_pct,
        metadata_changed=metadata_changed,
        structure_change_pct=structure_change_pct,
        file_size_change_pct=structure["file_size_change_pct"],
        classification=classification,
        signal_results=signal_results,
    )
    return {
        "stored_document_name": stored_record.get("document_name"),
        "stored_hash": stored_record.get("document_hash"),
        "stored_metadata": stored_meta,
        "uploaded_metadata": uploaded_meta,
        "metadata_changed": metadata_changed,
        "metadata_diff": metadata_diff,
        "content_change_pct": content_change_pct,
        "content_diff_length": diff_length,
        "content_total_length": total_length,
        "structure_change_pct": structure_change_pct,
        "file_size_change_pct": structure["file_size_change_pct"],
        "page_count_change_pct": structure["page_count_change_pct"],
        "risk_score": risk_score,
        "classification": classification,
        "similarity_pct": similarity_pct,
        "text_equivalent": text_equivalent,
        "signal_results": signal_results,
        "original_codes": original_codes,
        "uploaded_codes": uploaded_codes,
        "reasons": reasons,
        "diff_text": generate_text_diff(original_text, uploaded_text),
        "changed_sections": detect_text_changes(original_text, uploaded_text),
        "original_text_available": bool(original_text.strip()),
        "uploaded_text_available": bool(uploaded_text.strip()),
        "content_equivalent": content_equivalent,
        "metadata_status": "Original" if not metadata_changed or content_equivalent else "Changed",
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }


def login_required(required_role=None):
    def decorator(fn):
        @wraps(fn)
        def wrapper(*args, **kwargs):
            if not session.get("logged_in"):
                flash("Please login to continue.", "warning")
                return redirect(url_for("login"))
            if required_role and session.get("role") != required_role:
                flash("You are not authorized to access that page.", "danger")
                return redirect(url_for("dashboard"))
            return fn(*args, **kwargs)
        return wrapper
    return decorator


@app.context_processor
def inject_session():
    return {"session": session, "format_timestamp": format_timestamp}


# Ensure default users are available before any route requires login
ensure_user_db()


# ─────────────────────────────────────────────────────────────
# ROUTES — PAGES
# ─────────────────────────────────────────────────────────────

@app.route("/")
def index():
    """Home / Landing page."""
    stats = blockchain.get_stats()
    return render_template("index.html", stats=stats)


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        user = get_user(username)
        if user and verify_password(password, user["salt"], user["password_hash"]):
            session["logged_in"] = True
            session["username"] = username
            session["role"] = user["role"]
            append_audit_entry("login", "success", username=username, role=user["role"])
            flash(f"Welcome back, {username}!", "success")
            return redirect(url_for("dashboard"))
        append_audit_entry("login", "failed", username=username, details="invalid credentials")
        flash("Invalid username or password.", "danger")
    return render_template("login.html")


@app.route("/logout")
@login_required()
def logout():
    append_audit_entry("logout", "success")
    session.clear()
    flash("Logged out successfully.", "success")
    return redirect(url_for("login"))


@app.route("/dashboard")
@login_required()
def dashboard():
    stats = blockchain.get_stats()
    role = session.get("role")
    username = session.get("username")
    if role == "Admin":
        logs = get_audit_entries(20)
        documents = get_registered_documents_for_admin()
    else:
        logs = get_audit_entries(20, role="Verifier", username=username)
        documents = []
    return render_template("dashboard.html", stats=stats, logs=logs, documents=documents)


@app.route("/audit")
@login_required("Admin")
def audit():
    logs = get_audit_entries(100)
    audit_chain_valid = verify_audit_chain(load_audit_log())
    return render_template("audit.html", logs=logs, audit_chain_valid=audit_chain_valid)


@app.route("/documents/revoke", methods=["POST"])
@login_required("Admin")
def revoke_registered_document():
    doc_hash = request.form.get("document_hash", "").strip()
    if not doc_hash:
        flash("Missing document hash.", "danger")
        return redirect(url_for("dashboard"))
    revoke_document(doc_hash)
    append_audit_entry("delete_registered_document", "success", details=f"hash={doc_hash}")
    flash("Document removed from active registry.", "success")
    return redirect(url_for("dashboard"))


@app.route("/audit/delete", methods=["POST"])
@login_required()
def delete_audit_log_entry():
    entry_id = request.form.get("entry_id", "").strip()
    if not entry_id:
        flash("Missing log entry id.", "danger")
        return redirect(url_for("dashboard"))

    success = delete_audit_entry(
        entry_id=entry_id,
        requester_role=session.get("role", ""),
        requester_user=session.get("username", ""),
    )
    if success:
        flash("Log entry deleted.", "success")
    else:
        flash("Could not delete this log entry.", "warning")
    return redirect(url_for("dashboard"))


# ── REGISTER ─────────────────────────────────────────────────

@app.route("/register", methods=["GET", "POST"])
@login_required("Admin")
def register():
    """Phase 1: Document Registration."""
    if request.method == "GET":
        return render_template("register.html")

    file = request.files.get("document")
    if not file or file.filename == "":
        flash("No file selected.", "danger")
        append_audit_entry("register_document", "failed", details="no file selected")
        return redirect(url_for("register"))

    if not allowed_file(file.filename):
        flash("File type not allowed. Upload PDF, DOC, DOCX, PNG, JPG, or TXT.", "danger")
        append_audit_entry("register_document", "failed", details="unsupported file type")
        return redirect(url_for("register"))

    student_name = request.form.get("student_name", "").strip()
    degree = request.form.get("degree", "").strip()
    institution = request.form.get("institution", "").strip()
    document_category = request.form.get("document_category", "").strip() or "Academic"
    # AES encryption is mandatory for all registered originals.
    encrypt_doc = True

    if not (student_name and degree and institution):
        flash("Please fill all required fields.", "danger")
        append_audit_entry("register_document", "failed", details="missing form fields")
        return redirect(url_for("register"))

    filename = secure_filename(file.filename)
    original_path = os.path.join(ORIGINALS_DIR, filename)
    file.save(original_path)

    doc_hash = sha256_file(original_path)
    extracted_text = extract_text(original_path)
    content_hash = compute_text_hash_from_text(extracted_text)
    signature = sign_hash(doc_hash, private_key_pem)
    content_signature = sign_hash(content_hash, private_key_pem) if content_hash else ""
    certificate_signature = sign_with_certificate(doc_hash, private_key_pem)
    content_certificate_signature = sign_with_certificate(content_hash, private_key_pem) if content_hash else ""

    encrypted_file_path = ""
    with open(original_path, "rb") as f:
        raw = f.read()
    encrypted_payload = aes_encrypt(raw, aes_key)
    encrypted_file_path = os.path.join(ENCRYPTED_DIR, filename + ".enc")
    with open(encrypted_file_path, "w", encoding="utf-8") as f:
        json.dump(encrypted_payload, f)
    hmac_tag = generate_hmac(doc_hash, "INSTITUTION_SECRET_2024")
    content_hmac = generate_hmac(content_hash, "INSTITUTION_SECRET_2024") if content_hash else ""
    metadata = extract_metadata(original_path)

    record = build_document_record(
        filename=filename,
        doc_hash=doc_hash,
        student_name=student_name,
        degree=degree,
        institution=institution,
        signature=signature,
        public_key_pem=public_key_pem,
        metadata=metadata,
        original_file_path=original_path,
        encrypted_file_path=encrypted_file_path,
        encrypted=encrypt_doc,
        document_type=Path(filename).suffix.lower(),
        document_category=document_category,
        certificate_pem=institution_certificate_pem,
        certificate_signature=certificate_signature,
        content_certificate_signature=content_certificate_signature,
    )
    record["content_hash"] = content_hash
    record["content_signature"] = content_signature
    record["content_hmac"] = content_hmac
    record["text_hash_available"] = bool(content_hash)

    new_block = blockchain.add_document(record)
    restore_document(doc_hash)
    append_audit_entry("register_document", "success", document_name=filename)

    context = {
        "filename": filename,
        "student_name": student_name,
        "degree": degree,
        "institution": institution,
        "document_category": document_category,
        "doc_hash": doc_hash,
        "content_hash": content_hash,
        "signature": signature[:60] + "...",
        "certificate_subject": institution_certificate_info.get("subject"),
        "hmac_tag": hmac_tag,
        "block_index": new_block.index,
        "block_hash": new_block.hash,
        "nonce": new_block.nonce,
        "merkle_root": new_block.merkle_root,
        "timestamp": new_block.timestamp,
        "file_size": file_size_kb(original_path),
        "encrypted": encrypt_doc,
    }
    return render_template("register_result.html", **context)


# ── VERIFY ───────────────────────────────────────────────────

@app.route("/verify", methods=["GET", "POST"])
@login_required("Verifier")
def verify():
    """Phase 2: Document Verification."""
    registered_documents = get_registered_documents_for_verifier()
    if request.method == "GET":
        return render_template("verify.html", registered_documents=registered_documents)

    file = request.files.get("document")
    selected_document_hash = request.form.get("registered_document_hash", "").strip()
    if not file or file.filename == "":
        flash("No file selected.", "danger")
        append_audit_entry("verify_document", "failed", details="no file selected")
        return redirect(url_for("verify"))

    if not selected_document_hash:
        flash("Please select a registered original document first.", "warning")
        append_audit_entry("verify_document", "failed", details="no selected original document")
        return redirect(url_for("verify"))

    if not allowed_file(file.filename):
        flash("File type not allowed.", "danger")
        append_audit_entry("verify_document", "failed", details="unsupported file type")
        return redirect(url_for("verify"))

    filename = secure_filename(file.filename)
    temp_path = os.path.join(app.config["UPLOAD_FOLDER"], f"verify_{datetime.now(timezone.utc).timestamp()}_{filename}")
    file.save(temp_path)
    uploaded_hash = sha256_file(temp_path)
    uploaded_content_hash = compute_text_hash_from_file(temp_path)
    report_timestamp = datetime.now(timezone.utc).isoformat()

    selected_block, selected_tx = blockchain.find_document(selected_document_hash)
    result = {
        "filename": filename,
        "selected_document_hash": selected_document_hash,
        "selected_document_name": selected_tx.get("document_name") if selected_tx else None,
        "uploaded_hash": uploaded_hash,
        "uploaded_content_hash": uploaded_content_hash,
        "file_size": file_size_kb(temp_path),
        "analysis": None,
        "status": None,
        "verification_timestamp": report_timestamp,
        "verification_report": None,
    }

    if not selected_tx or not is_document_active(selected_document_hash):
        result.update({
            "verified": False,
            "analysis": None,
            "status": "Not Registered",
        })
        result["verification_report"] = {
            "document_name": filename,
            "status": "Not Registered",
            "risk_score": 100.0,
            "metadata_status": "Unknown",
            "timestamp": report_timestamp,
        }
        append_audit_entry("verify_document", "failed", document_name=filename, details="selected original not found or inactive")
    else:
        stored_content_hash = selected_tx.get("content_hash")
        if not stored_content_hash:
            stored_content_hash = compute_text_hash_from_file(selected_tx.get("original_file_path", ""))

        file_hash_match = uploaded_hash == selected_document_hash
        content_hash_match = bool(uploaded_content_hash and stored_content_hash and uploaded_content_hash == stored_content_hash)
        result["stored_content_hash"] = stored_content_hash
        result["content_hash_matched"] = content_hash_match

        if file_hash_match:
            sig_valid = verify_rsa_signature(uploaded_hash, selected_tx["signature"], selected_tx["public_key"])
            hmac_valid = verify_hmac(uploaded_hash, "INSTITUTION_SECRET_2024", selected_tx["hmac"])
            cert_valid = verify_certificate_signature(
                uploaded_hash,
                selected_tx.get("certificate_signature", ""),
                selected_tx.get("certificate_pem", institution_certificate_pem),
            )
            match_basis = "binary_hash"
        elif content_hash_match:
            content_signature = selected_tx.get("content_signature")
            content_hmac = selected_tx.get("content_hmac")
            sig_valid = verify_rsa_signature(uploaded_content_hash, content_signature, selected_tx["public_key"]) if content_signature else True
            hmac_valid = verify_hmac(uploaded_content_hash, "INSTITUTION_SECRET_2024", content_hmac) if content_hmac else True
            cert_valid = verify_certificate_signature(
                uploaded_content_hash,
                selected_tx.get("content_certificate_signature", ""),
                selected_tx.get("certificate_pem", institution_certificate_pem),
            ) if selected_tx.get("content_certificate_signature") else True
            match_basis = "text_content_hash"
        else:
            sig_valid = False
            hmac_valid = False
            cert_valid = False
            match_basis = "none"

        if file_hash_match or content_hash_match:
            result.update({
                "verified": True,
                "student_name": selected_tx["student_name"],
                "degree": selected_tx["degree"],
                "institution": selected_tx["institution"],
                "issued_at": selected_tx["issued_at"],
                "signature_valid": sig_valid,
                "hmac_valid": hmac_valid,
                "certificate_valid": cert_valid,
                "block_index": selected_block["index"],
                "block_hash": selected_block["hash"],
                "merkle_root": selected_block["merkle_root"],
                "nonce": selected_block["nonce"],
                "stored_hash": selected_tx["document_hash"],
                "encrypted": selected_tx.get("encrypted", False),
                "status": "Original",
                "match_basis": match_basis,
            })
            result["verification_report"] = {
                "document_name": filename,
                "status": "Original",
                "risk_score": 0.0,
                "metadata_status": "Original",
                "timestamp": report_timestamp,
            }
            append_audit_entry("verify_document", "original", document_name=filename, details=f"basis={match_basis}")
        else:
            # Explicit selected document comparison for tamper detection.
            analysis = analyze_uploaded_against_stored(selected_tx, temp_path)
            if analysis["content_equivalent"]:
                result.update({
                    "verified": True,
                    "student_name": selected_tx["student_name"],
                    "degree": selected_tx["degree"],
                    "institution": selected_tx["institution"],
                    "issued_at": selected_tx["issued_at"],
                    "signature_valid": True,
                    "hmac_valid": True,
                    "certificate_valid": True,
                    "block_index": selected_block["index"],
                    "block_hash": selected_block["hash"],
                    "merkle_root": selected_block["merkle_root"],
                    "nonce": selected_block["nonce"],
                    "stored_hash": selected_tx["document_hash"],
                    "encrypted": selected_tx.get("encrypted", False),
                    "status": "Original",
                    "match_basis": "semantic_text_match",
                    "analysis": analysis,
                })
                result["verification_report"] = {
                    "document_name": filename,
                    "status": "Original",
                    "risk_score": 0.0,
                    "metadata_status": "Original",
                    "timestamp": report_timestamp,
                }
                append_audit_entry("verify_document", "original", document_name=filename, details="basis=semantic_text_match")
            else:
                cert_binary_valid = verify_certificate_signature(
                    uploaded_hash,
                    selected_tx.get("certificate_signature", ""),
                    selected_tx.get("certificate_pem", institution_certificate_pem),
                ) if selected_tx.get("certificate_signature") else False
                cert_content_valid = verify_certificate_signature(
                    uploaded_content_hash,
                    selected_tx.get("content_certificate_signature", ""),
                    selected_tx.get("certificate_pem", institution_certificate_pem),
                ) if uploaded_content_hash and selected_tx.get("content_certificate_signature") else False
                raw_hmac_valid = verify_hmac(uploaded_hash, "INSTITUTION_SECRET_2024", selected_tx.get("hmac", "")) if selected_tx.get("hmac") else False
                content_hmac_valid = verify_hmac(uploaded_content_hash, "INSTITUTION_SECRET_2024", selected_tx.get("content_hmac", "")) if uploaded_content_hash and selected_tx.get("content_hmac") else False
                analysis["signal_results"].extend([
                    build_signal(
                        "PKI certificate",
                        0.0 if (cert_binary_valid or cert_content_valid) else 100.0,
                        0.15,
                        "matched" if (cert_binary_valid or cert_content_valid) else "mismatch",
                        "Certificate-backed signature validates." if (cert_binary_valid or cert_content_valid) else "Certificate-backed signature does not validate against the uploaded document.",
                    ),
                    build_signal(
                        "HMAC integrity",
                        0.0 if (raw_hmac_valid or content_hmac_valid) else 100.0,
                        0.10,
                        "matched" if (raw_hmac_valid or content_hmac_valid) else "mismatch",
                        "HMAC integrity check validates." if (raw_hmac_valid or content_hmac_valid) else "HMAC integrity check does not validate.",
                    ),
                ])
                analysis["risk_score"], analysis["classification"] = compute_ensemble_verdict(analysis["signal_results"])
                if analysis["classification"] == "Safe" and analysis["content_change_pct"] > 0:
                    analysis["classification"] = "Suspicious"
                    analysis["risk_score"] = max(analysis["risk_score"], 25.0)
                analysis["reasons"] = build_analysis_reasons(
                    content_equivalent=analysis["content_equivalent"],
                    original_text_available=analysis["original_text_available"],
                    uploaded_text_available=analysis["uploaded_text_available"],
                    content_change_pct=analysis["content_change_pct"],
                    metadata_changed=analysis["metadata_changed"],
                    structure_change_pct=analysis["structure_change_pct"],
                    file_size_change_pct=analysis["file_size_change_pct"],
                    classification=analysis["classification"],
                    signal_results=analysis["signal_results"],
                )
                computed_status = analysis["classification"]
                result.update({
                    "verified": False,
                    "analysis": analysis,
                    "status": computed_status,
                    "candidate_name": selected_tx.get("document_name"),
                    "stored_hash": selected_tx.get("document_hash"),
                    "match_basis": "none",
                    "certificate_valid": cert_binary_valid or cert_content_valid,
                })
                result["verification_report"] = {
                    "document_name": filename,
                    "status": computed_status,
                    "risk_score": analysis["risk_score"],
                    "metadata_status": analysis["metadata_status"],
                    "timestamp": report_timestamp,
                }
                append_audit_entry("verify_document", "tampered", document_name=filename, details=f"risk={analysis['risk_score']}")

    try:
        os.remove(temp_path)
    except OSError:
        pass

    return render_template("verify_result.html", **result)


# ── BLOCKCHAIN EXPLORER ───────────────────────────────────────

@app.route("/explorer")
def explorer():
    """Blockchain Explorer — view all blocks."""
    chain    = blockchain.to_dict()
    is_valid, messages = blockchain.is_chain_valid()
    stats    = blockchain.get_stats()
    return render_template("explorer.html",
                           chain=chain,
                           is_valid=is_valid,
                           messages=messages,
                           stats=stats)


# ── DEMO: HASH CALCULATOR ─────────────────────────────────────

@app.route("/demo/hash", methods=["GET", "POST"])
def demo_hash():
    """Interactive SHA-256 demo — type text, see hash change live."""
    result = None
    if request.method == "POST":
        text = request.form.get("text", "")
        result = {
            "input":   text,
            "sha256":  sha256_string(text),
            "length":  len(text),
        }
    return render_template("demo_hash.html", result=result)


# ── DEMO: TAMPER SIMULATOR ────────────────────────────────────

@app.route("/demo/tamper", methods=["GET", "POST"])
def demo_tamper():
    """Side-by-side hash comparison to show avalanche effect."""
    result = None
    if request.method == "POST":
        text1 = request.form.get("text1", "")
        text2 = request.form.get("text2", "")
        h1    = sha256_string(text1)
        h2    = sha256_string(text2)
        # Count differing hex characters
        diff_chars = sum(1 for a, b in zip(h1, h2) if a != b)
        result = {
            "text1":       text1,
            "text2":       text2,
            "hash1":       h1,
            "hash2":       h2,
            "match":       h1 == h2,
            "diff_chars":  diff_chars,
            "diff_percent": round(diff_chars / 64 * 100, 1),
        }
    return render_template("demo_tamper.html", result=result)


# ── DEMO: ENCRYPTION ─────────────────────────────────────────

@app.route("/demo/encrypt", methods=["GET", "POST"])
def demo_encrypt():
    """AES-256 encryption demo."""
    result = None
    if request.method == "POST":
        plaintext = request.form.get("plaintext", "")
        enc = aes_encrypt(plaintext.encode(), aes_key)
        decrypted = aes_decrypt(enc["ciphertext"], enc["iv"], aes_key).decode(errors="replace")
        result = {
            "plaintext":   plaintext,
            "ciphertext":  enc["ciphertext"][:80] + "..." if len(enc["ciphertext"]) > 80 else enc["ciphertext"],
            "iv":          enc["iv"],
            "decrypted":   decrypted,
            "key_preview": base64.b64encode(aes_key).decode()[:20] + "...",
        }
    return render_template("demo_encrypt.html", result=result)


# ── API ENDPOINTS ─────────────────────────────────────────────

@app.route("/api/blockchain")
def api_blockchain():
    """Return full blockchain as JSON."""
    return jsonify(blockchain.to_dict())


@app.route("/api/validate")
def api_validate():
    """Validate blockchain integrity."""
    valid, messages = blockchain.is_chain_valid()
    return jsonify({"valid": valid, "messages": messages})


@app.route("/api/stats")
def api_stats():
    return jsonify(blockchain.get_stats())


@app.route("/api/search_hash", methods=["POST"])
def api_search_hash():
    """Search blockchain by document hash (for manual entry)."""
    data = request.get_json(silent=True) or {}
    doc_hash = data.get("hash", "").strip()
    if not doc_hash:
        return jsonify({"error": "No hash provided"}), 400
    block, tx = blockchain.find_document(doc_hash)
    if block:
        return jsonify({"found": True, "block": block, "transaction": tx})
    return jsonify({"found": False})


# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("\n" + "="*60)
    print("  SECURE ACADEMIC DOCUMENT VERIFICATION SYSTEM")
    print("="*60)
    print(f"  Blockchain blocks : {len(blockchain.chain)}")
    print(f"  PoW Difficulty    : {blockchain.difficulty}")
    print(f"  RSA Key           : 2048-bit loaded")
    print(f"  AES Key           : 256-bit loaded")
    print("="*60)
    print("  Open http://127.0.0.1:5000")
    print("="*60 + "\n")
    app.run(debug=True, port=5000)
